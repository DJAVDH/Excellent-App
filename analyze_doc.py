"""Analyze Word document structure"""
from docx import Document
import os

try:
    doc = Document("Werkbon Assistentie.docx")
    print("✓ Document loaded successfully")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    
    # Check for images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
            print(f"Image found: {rel.target_ref}")
    
    print(f"Total images: {image_count}")
    
    # Show first few paragraphs
    print("\nContent preview:")
    for i, para in enumerate(doc.paragraphs[:15]):
        if para.text.strip():
            print(f"  P{i}: {para.text}")
    
    # Show table structure
    if doc.tables:
        print(f"\nTable structure:")
        for ti, table in enumerate(doc.tables):
            print(f"  Table {ti}: {len(table.rows)} rows x {len(table.columns)} cols")
            print(f"    Columns: {len(table.columns)}")
            if table.rows:
                for ci, cell in enumerate(table.rows[0].cells[:10]):
                    print(f"      Col {ci}: {cell.text[:40]}")
                print(f"    Row 1 data: {[cell.text[:20] for cell in table.rows[1].cells[:5]]}")
    
except FileNotFoundError as e:
    print(f"File not found: {e}")
    print(f"Current dir: {os.getcwd()}")
    print(f"Files: {os.listdir('.')}")
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
